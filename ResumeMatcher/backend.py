# backend.py
"""
Backend logic for RAG Resume Matcher:
- text extraction from upload object
- chunking & embedding resume into FAISS
- retrieval of top-k resume snippets for JD
- craft prompt for Ollama with 8 evaluation criteria
- call Ollama LLM and parse JSON output
"""

import io
import re
import json
from typing import List
from dataclasses import dataclass

import pdfplumber
import docx
import numpy as np

# LangChain imports (attempt both main and community paths)
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.docstore.document import Document
    from langchain.vectorstores import FAISS
    try:
        from langchain.embeddings import OllamaEmbeddings
    except Exception:
        from langchain_community.embeddings import OllamaEmbeddings

    try:
        from langchain.llms import Ollama
    except Exception:
        from langchain_community.llms import Ollama
except Exception as e:
    raise ImportError("LangChain or Ollama imports failed. Check installed packages.") from e


@dataclass
class MatchResult:
    candidate_name: str
    fit_evaluation: dict
    overall_fit_score: float
    recommendation: str
    summary: str
    raw_llm_output: str
    retrieved: List[dict]


# ---------- Helpers ----------
def extract_text_from_upload(uploaded_file) -> str:
    """Extract text from an uploaded file object (Streamlit file_uploader)."""
    if not uploaded_file:
        return ""
    name = uploaded_file.name.lower()
    data = uploaded_file.read()
    if isinstance(data, str):
        data = data.encode("utf-8")

    if name.endswith(".txt"):
        return data.decode(errors="ignore")
    if name.endswith(".pdf"):
        pages = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for p in pdf.pages:
                txt = p.extract_text()
                if txt:
                    pages.append(txt)
        return "\n".join(pages)
    if name.endswith(".docx") or name.endswith(".doc"):
        doc = docx.Document(io.BytesIO(data))
        return "\n".join([p.text for p in doc.paragraphs])
    try:
        return data.decode(errors="ignore")
    except Exception:
        return ""


def simple_token_set(text: str):
    text = re.sub(r"[^\w\s+-]", " ", (text or "").lower())
    tokens = [t for t in text.split() if len(t) > 2]
    stop = {"the", "and", "for", "with", "that", "this", "from", "your", "you", "our", "a", "an", "to", "in", "on",
            "of", "is", "are"}
    return set([t for t in tokens if t not in stop])


# ---------- RAG functions ----------
def build_resume_vectorstore(resume_text: str, embed_model: str, chunk_size=800, chunk_overlap=150):
    """Chunk resume and build a FAISS vectorstore using OllamaEmbeddings."""
    if not resume_text or not resume_text.strip():
        return None, []
    embeddings = OllamaEmbeddings(model=embed_model)
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    chunks = splitter.split_text(resume_text)
    docs = [Document(page_content=c, metadata={"chunk_id": i}) for i, c in enumerate(chunks)]
    if len(docs) == 0:
        return None, []
    vectorstore = FAISS.from_documents(docs, embeddings)
    return vectorstore, chunks


def retrieve_for_jd(vectorstore, jd_text: str, top_k=6):
    """Retrieve top_k resume chunks for the given JD text."""
    if vectorstore is None:
        return []
    retriever = vectorstore.as_retriever(search_kwargs={"k": top_k})
    docs = retriever.get_relevant_documents(jd_text)
    return [{"text": d.page_content, "metadata": getattr(d, "metadata", {})} for d in docs]


# ---------- Prompt & parsing ----------
def craft_detailed_prompt(jd: str, retrieved_chunks: List[str]) -> str:
    """Craft a strict prompt asking for the 8-criteria JSON output (only JSON)."""
    instructions = r"""
You are an expert technical recruiter. Evaluate the candidate using ONLY the explicitly provided resume snippets.
Use the following 8 criteria and return a single JSON object (no extra text) with exact keys described in the Output Format.

1. Skill Match - score 0.0-1.0 - list overlapping skills and quotes from resume.
2. Job Switch Frequency - score 0.0-1.0 - extract employment durations and count short jobs (<1.5 yrs).
3. Education Completeness - score 0.0-1.0 - check for degree, institution, year.
4. Work Experience Completeness - score 0.0-1.0 - identify gaps >6 months, missing title/company/dates.
5. Career Progression - score 0.0-1.0 - detect role growth over time.
6. Domain Relevance - score 0.0-1.0 - alignment with target domain mentioned in the JD.
7. Certifications - score 0.0-1.0 - list certifications verbatim.
8. Soft Skills & Leadership - score 0.0-1.0 - quote lines showing leadership/teamwork/communication.

Be objective and use only explicitly stated information. If a field cannot be determined, set score to 0.0 and leave reference list empty.

Output Format (JSON EXACTLY):
{
  "candidate_name": "string or empty",
  "fit_evaluation": {
    "skill_match": {"score": 0.00, "reference": []},
    "job_switch_frequency": {"score": 0.00, "reference": []},
    "education_completeness": {"score": 0.00, "reference": []},
    "experience_completeness": {"score": 0.00, "reference": []},
    "career_progression": {"score": 0.00, "reference": []},
    "domain_relevance": {"score": 0.00, "reference": []},
    "certifications": {"score": 0.00, "reference": []},
    "soft_skills": {"score": 0.00, "reference": []}
  },
  "overall_fit_score": 0.00,
  "recommendation": "Shortlist / Consider / Reject",
  "summary": "concise reasoning (<=40 words)"
}
RETURN ONLY THE JSON OBJECT - NO ADDITIONAL TEXT.
"""

    snippets_text = ""
    for i, s in enumerate(retrieved_chunks, start=1):
        # provide each snippet clearly
        snippets_text += f"\n\nSnippet {i}:\n{s}"

    return f"{instructions}\n\nJob Description:\n{jd}\n\nRelevant Resume Snippets:\n{snippets_text}\n\n"


def extract_json_from_text(text: str):
    """Find and parse a JSON object from the model output robustly."""
    # find the first "{" and last matching "}" and attempt json.loads
    start = text.find("{")
    if start == -1:
        raise ValueError("No JSON object start brace found.")
    # try to find matching end by scanning forward (handles nested braces)
    depth = 0
    end_idx = -1
    for i in range(start, len(text)):
        if text[i] == "{":
            depth += 1
        elif text[i] == "}":
            depth -= 1
            if depth == 0:
                end_idx = i
                break
    if end_idx == -1:
        raise ValueError("Couldn't find end of JSON object.")
    json_str = text[start:end_idx + 1]
    return json.loads(json_str)


# ---------- Main orchestrator ----------
def run_resume_match(
    jd_text: str,
    resume_text: str,
    llm_model: str = "llama2",
    embed_model: str = "nomic-embed-text",
    chunk_size: int = 800,
    chunk_overlap: int = 150,
    top_k: int = 6,
) -> MatchResult:
    # 1) Build vector store from resume
    vectorstore, chunks = build_resume_vectorstore(resume_text, embed_model, chunk_size, chunk_overlap)
    if vectorstore is None:
        # return blank result
        return MatchResult(candidate_name="",
                           fit_evaluation={},
                           overall_fit_score=0.0,
                           recommendation="Reject",
                           summary="No resume content provided or chunking failed.",
                           raw_llm_output="",
                           retrieved=[])

    # 2) Retrieve
    retrieved_docs = retrieve_for_jd(vectorstore, jd_text, top_k)
    retrieved_texts = [d["text"] for d in retrieved_docs]

    # 3) Build prompt
    prompt = craft_detailed_prompt(jd_text, retrieved_texts)

    # 4) Call Ollama LLM
    llm = Ollama(model=llm_model)
    llm_response = llm(prompt)  # wrapper may return text or object; cast to str
    llm_text = llm_response.content if hasattr(llm_response, "content") else str(llm_response)

    # 5) Parse JSON
    try:
        parsed = extract_json_from_text(llm_text)
    except Exception as e:
        # If parsing fails, return raw LLM output in summary and a Reject
        return MatchResult(candidate_name="",
                           fit_evaluation={},
                           overall_fit_score=0.0,
                           recommendation="Reject",
                           summary=f"LLM output not parseable: {str(e)}",
                           raw_llm_output=llm_text,
                           retrieved=retrieved_docs)

    # 6) Normalize and return
    candidate_name = parsed.get("candidate_name", "")
    fit_evaluation = parsed.get("fit_evaluation", {})
    overall_fit_score = float(parsed.get("overall_fit_score", 0.0))
    # Normalize to 0–5 scale if value seems to be 0–1
    if overall_fit_score <= 1.0:
        overall_fit_score = round(overall_fit_score * 5, 2)

    # Adjust recommendation based on normalized score
    if overall_fit_score < 2.5:
        recommendation = "Reject"
    elif 2.5 <= overall_fit_score < 3.5:
        recommendation = "Consider"
    else:
        recommendation = "Shortlist"

    # recommendation = parsed.get("recommendation", "Consider")
    summary = parsed.get("summary", "")

    return MatchResult(candidate_name=candidate_name,
                       fit_evaluation=fit_evaluation,
                       overall_fit_score=overall_fit_score,
                       recommendation=recommendation,
                       summary=summary,
                       raw_llm_output=llm_text,
                       retrieved=retrieved_docs)
